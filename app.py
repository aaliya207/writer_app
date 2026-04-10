from flask import Flask, render_template, request, jsonify, send_file, session, redirect, url_for, abort
from flask_sqlalchemy import SQLAlchemy
from datetime import datetime, timedelta
from config import Config
import os, urllib.parse, requests as http_requests, secrets, re, threading, io, json, base64
from googleapiclient.discovery import build
from googleapiclient.http import MediaInMemoryUpload
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, HRFlowable
from reportlab.lib import colors
from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup
import json
import sys
if sys.stdout is not None:
    sys.stdout.reconfigure(encoding='utf-8')
# Handle PyInstaller bundle paths
if hasattr(sys, '_MEIPASS'):
    BASE_DIR = sys._MEIPASS
else:
    BASE_DIR = os.path.abspath('.')

app = Flask(__name__,
    template_folder=os.path.join(BASE_DIR, 'templates'),
    static_folder=os.path.join(BASE_DIR, 'static')
)

app = Flask(__name__)
app.config.from_object(Config)
# Store database in user's home directory so it persists
USER_DATA_DIR = os.path.join(os.path.expanduser('~'), 'Scripvia')
os.makedirs(USER_DATA_DIR, exist_ok=True)
app.config['SQLALCHEMY_DATABASE_URI'] = f"sqlite:///{os.path.join(USER_DATA_DIR, 'scripvia.db')}"
app.permanent_session_lifetime = timedelta(days=30)

SESSION_FILE = os.path.join(USER_DATA_DIR, 'session.json')

def save_persistent_session(user_id=None, is_guest=False, guest_session_id=None):
    try:
        with open(SESSION_FILE, 'w') as f:
            json.dump({
                'user_id': user_id,
                'is_guest': is_guest,
                'guest_session_id': guest_session_id
            }, f)
    except Exception as e:
        print(f"Session save error: {e}")

def load_persistent_session():
    try:
        if os.path.exists(SESSION_FILE):
            with open(SESSION_FILE, 'r') as f:
                return json.load(f)
    except Exception:
        pass
    return None

def clear_persistent_session():
    try:
        if os.path.exists(SESSION_FILE):
            os.remove(SESSION_FILE)
    except Exception:
        pass

db = SQLAlchemy(app)
from flask_migrate import Migrate
migrate = Migrate(app, db)

os.environ['OAUTHLIB_INSECURE_TRANSPORT'] = '1'


# --- MODELS ---

class User(db.Model):
    id                 = db.Column(db.Integer, primary_key=True)
    google_id          = db.Column(db.String(200), unique=True, nullable=False)
    email              = db.Column(db.String(200), nullable=False)
    name               = db.Column(db.String(200), default='')
    picture            = db.Column(db.String(500), default='')
    access_token       = db.Column(db.Text, default='')
    refresh_token      = db.Column(db.Text, default='')
    scripvia_folder_id = db.Column(db.String(200), default='')
    created_at         = db.Column(db.DateTime, default=datetime.utcnow)
    projects           = db.relationship('Project', backref='owner', lazy=True)

    def to_dict(self):
        return {'id': self.id, 'email': self.email, 'name': self.name, 'picture': self.picture}


CREATIVE_GENRES = ['fantasy', 'sci-fi', 'fiction', 'romance', 'mystery', 'thriller', 'horror', 'historical']

class Project(db.Model):
    id              = db.Column(db.Integer, primary_key=True)
    title           = db.Column(db.String(200), nullable=False)
    description     = db.Column(db.Text, default='')
    genre           = db.Column(db.String(100), default='general')
    map_image_url   = db.Column(db.Text, default='')
    user_id         = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=True)
    guest_session_id = db.Column(db.String(200), default='')
    drive_folder_id = db.Column(db.String(200), default='')
    created_at      = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at      = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)
    documents       = db.relationship('Document',  backref='project', lazy=True, cascade='all, delete-orphan')
    characters      = db.relationship('Character', backref='project', lazy=True, cascade='all, delete-orphan')
    scenes          = db.relationship('Scene',     backref='project', lazy=True, cascade='all, delete-orphan')
    lore_items      = db.relationship('LoreItem',  backref='project', lazy=True, cascade='all, delete-orphan')
    notes           = db.relationship('Note',      backref='project', lazy=True, cascade='all, delete-orphan')
    lore_relationships = db.relationship('LoreRelationship', backref='project', lazy=True, cascade='all, delete-orphan')
    CREATIVE_GENRES = CREATIVE_GENRES

    def to_dict(self):
        return {
            'id': self.id, 'title': self.title, 'description': self.description,
            'map_image_url': self.map_image_url or '',
            'genre': self.genre, 'is_creative': self.genre in CREATIVE_GENRES,
            'created_at': self.created_at.isoformat(), 'updated_at': self.updated_at.isoformat(),
            'document_count': len(self.documents)
        }


class Document(db.Model):
    id            = db.Column(db.Integer, primary_key=True)
    title         = db.Column(db.String(200), nullable=False)
    content       = db.Column(db.Text, default='')
    project_id    = db.Column(db.Integer, db.ForeignKey('project.id'), nullable=False)
    drive_file_id = db.Column(db.String(200), default='')
    order_index   = db.Column(db.Integer, default=0)
    created_at    = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at    = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)

    def to_dict(self):
        return {
            'id': self.id, 'title': self.title, 'content': self.content,
            'project_id': self.project_id, 'drive_file_id': self.drive_file_id,
            'created_at': self.created_at.isoformat(), 'updated_at': self.updated_at.isoformat()
        }

class Character(db.Model):
    id          = db.Column(db.Integer, primary_key=True)
    project_id  = db.Column(db.Integer, db.ForeignKey('project.id'), nullable=False)
    name        = db.Column(db.String(200), nullable=False)
    role        = db.Column(db.String(100), default='')
    age         = db.Column(db.String(50), default='')
    personality = db.Column(db.Text, default='')
    backstory   = db.Column(db.Text, default='')
    appearance  = db.Column(db.Text, default='')
    titles      = db.Column(db.Text, default='[]')
    image_url    = db.Column(db.String(500), default='')
    image_focus  = db.Column(db.String(20), default='center')
    extra_notes  = db.Column(db.Text, default='')
    drive_file_id = db.Column(db.String(200), default='')
    drive_image_file_id = db.Column(db.String(200), default='')
    web_x       = db.Column(db.Float, nullable=True)
    web_y       = db.Column(db.Float, nullable=True)
    created_at  = db.Column(db.DateTime, default=datetime.utcnow)

    def to_dict(self):
        return {
            'id': self.id, 'project_id': self.project_id, 'name': self.name,
            'role': self.role, 'age': self.age, 'personality': self.personality,
            'backstory': self.backstory, 'appearance': self.appearance, 'titles': self.get_titles(),
            'image_url': self.image_url, 'image_focus': self.image_focus, 'extra_notes': self.extra_notes,
            'drive_file_id': self.drive_file_id, 'drive_image_file_id': self.drive_image_file_id,
            'web_x': self.web_x, 'web_y': self.web_y,
            'created_at': self.created_at.isoformat()
        }

    def get_titles(self):
        try:
            parsed = json.loads(self.titles or '[]')
            return [str(title).strip() for title in parsed if str(title).strip()]
        except Exception:
            return []

    def set_titles(self, titles):
        clean_titles = [str(title).strip() for title in (titles or []) if str(title).strip()]
        self.titles = json.dumps(clean_titles)


class Scene(db.Model):
    id                = db.Column(db.Integer, primary_key=True)
    project_id        = db.Column(db.Integer, db.ForeignKey('project.id'), nullable=False)
    title             = db.Column(db.String(200), nullable=False)
    content           = db.Column(db.Text, default='')
    mood              = db.Column(db.String(100), default='')
    connected_chapter = db.Column(db.String(200), default='')
    drive_file_id     = db.Column(db.String(200), default='')
    created_at        = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at        = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)

    def to_dict(self):
        return {
            'id': self.id, 'project_id': self.project_id, 'title': self.title,
            'content': self.content, 'mood': self.mood,
            'connected_chapter': self.connected_chapter,
            'drive_file_id': self.drive_file_id,
            'created_at': self.created_at.isoformat(), 'updated_at': self.updated_at.isoformat()
        }


class LoreItem(db.Model):
    id          = db.Column(db.Integer, primary_key=True)
    project_id  = db.Column(db.Integer, db.ForeignKey('project.id'), nullable=False)
    name        = db.Column(db.String(200), nullable=False)
    category    = db.Column(db.String(100), default='item')
    description = db.Column(db.Text, default='')
    aliases     = db.Column(db.Text, default='[]')
    image_url   = db.Column(db.String(500), default='')
    image_focus = db.Column(db.String(20), default='center')
    extra_notes = db.Column(db.Text, default='')
    event_date  = db.Column(db.String(120), default='')
    event_order = db.Column(db.Integer, default=0)
    show_in_web = db.Column(db.Boolean, default=True)
    web_x       = db.Column(db.Float, nullable=True)
    web_y       = db.Column(db.Float, nullable=True)
    map_x       = db.Column(db.Float, default=50.0)
    map_y       = db.Column(db.Float, default=50.0)
    drive_file_id = db.Column(db.String(200), default='')
    drive_image_file_id = db.Column(db.String(200), default='')
    created_at  = db.Column(db.DateTime, default=datetime.utcnow)

    def to_dict(self):
        return {
            'id': self.id, 'project_id': self.project_id, 'name': self.name,
            'category': self.category, 'description': self.description, 'aliases': self.get_aliases(),
            'image_url': self.image_url, 'image_focus': self.image_focus, 'extra_notes': self.extra_notes,
            'event_date': self.event_date, 'event_order': self.event_order, 'show_in_web': bool(self.show_in_web),
            'web_x': self.web_x, 'web_y': self.web_y,
            'map_x': self.map_x, 'map_y': self.map_y,
            'drive_file_id': self.drive_file_id, 'drive_image_file_id': self.drive_image_file_id,
            'created_at': self.created_at.isoformat()
        }

    def get_aliases(self):
        try:
            parsed = json.loads(self.aliases or '[]')
            return [str(alias).strip() for alias in parsed if str(alias).strip()]
        except Exception:
            return []

    def set_aliases(self, aliases):
        clean_aliases = [str(alias).strip() for alias in (aliases or []) if str(alias).strip()]
        self.aliases = json.dumps(clean_aliases)


class Note(db.Model):
    id         = db.Column(db.Integer, primary_key=True)
    project_id = db.Column(db.Integer, db.ForeignKey('project.id'), nullable=False)
    content    = db.Column(db.Text, default='')
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)

    def to_dict(self):
        return {'id': self.id, 'project_id': self.project_id,
                'content': self.content, 'updated_at': self.updated_at.isoformat()}


class CharacterRelationship(db.Model):
    id            = db.Column(db.Integer, primary_key=True)
    project_id    = db.Column(db.Integer, db.ForeignKey('project.id'), nullable=False)
    char_a_id     = db.Column(db.Integer, db.ForeignKey('character.id'), nullable=False)
    char_b_id     = db.Column(db.Integer, db.ForeignKey('character.id'), nullable=False)
    relation_type = db.Column(db.String(100), default='')
    description   = db.Column(db.Text, default='')
    color         = db.Column(db.String(20), default='#7b6fb0')
    char_a        = db.relationship('Character', foreign_keys=[char_a_id], overlaps="char_b")
    char_b        = db.relationship('Character', foreign_keys=[char_b_id], overlaps="char_a")

    def to_dict(self):
        return {
            'id': self.id, 'project_id': self.project_id,
            'char_a_id': self.char_a_id, 'char_b_id': self.char_b_id,
            'char_a_name': self.char_a.name if self.char_a else '',
            'char_b_name': self.char_b.name if self.char_b else '',
            'relation_type': self.relation_type, 'description': self.description, 'color': self.color
        }


class LoreRelationship(db.Model):
    id            = db.Column(db.Integer, primary_key=True)
    project_id    = db.Column(db.Integer, db.ForeignKey('project.id'), nullable=False)
    lore_a_id     = db.Column(db.Integer, db.ForeignKey('lore_item.id'), nullable=False)
    lore_b_id     = db.Column(db.Integer, db.ForeignKey('lore_item.id'), nullable=False)
    relation_type = db.Column(db.String(100), default='')
    description   = db.Column(db.Text, default='')
    color         = db.Column(db.String(20), default='#5f8fd6')
    lore_a        = db.relationship('LoreItem', foreign_keys=[lore_a_id], overlaps="lore_b")
    lore_b        = db.relationship('LoreItem', foreign_keys=[lore_b_id], overlaps="lore_a")

    def to_dict(self):
        return {
            'id': self.id,
            'project_id': self.project_id,
            'lore_a_id': self.lore_a_id,
            'lore_b_id': self.lore_b_id,
            'lore_a_name': self.lore_a.name if self.lore_a else '',
            'lore_b_name': self.lore_b.name if self.lore_b else '',
            'relation_type': self.relation_type,
            'description': self.description,
            'color': self.color
        }


# --- HELPERS ---

def get_current_user():
    user_id = session.get('user_id')
    if not user_id:
        saved = load_persistent_session()
        if saved and saved.get('user_id'):
            session['user_id'] = saved['user_id']
            session.permanent = True
            user_id = saved['user_id']
    return User.query.get(user_id) if user_id else None

def get_guest_session_id():
    guest_session_id = session.get('guest_session_id')
    if not guest_session_id:
        saved = load_persistent_session()
        if saved and saved.get('is_guest') and saved.get('guest_session_id'):
            guest_session_id = saved['guest_session_id']
        else:
            guest_session_id = secrets.token_urlsafe(24)
        session['guest_session_id'] = guest_session_id
        session.permanent = True
    return guest_session_id


def legacy_project_filter():
    return Project.user_id.is_(None) & ((Project.guest_session_id == '') | Project.guest_session_id.is_(None))


def is_legacy_project(project):
    return project.user_id is None and not (project.guest_session_id or '').strip()


def project_is_owned_by_current_session(project):
    if is_legacy_project(project):
        return True
    user = get_current_user()
    if user:
        return project.user_id == user.id
    return project.user_id is None and project.guest_session_id == get_guest_session_id()


def claim_legacy_project(project):
    if not is_legacy_project(project):
        return False
    user = get_current_user()
    if user:
        project.user_id = user.id
        project.guest_session_id = ''
    else:
        project.user_id = None
        project.guest_session_id = get_guest_session_id()
    db.session.commit()
    return True


def get_project_or_404(project_id):
    project = Project.query.get_or_404(project_id)
    if not project_is_owned_by_current_session(project):
        abort(404)
    if request.method != 'GET':
        claim_legacy_project(project)
    return project


def get_drive_service(user):
    from google.oauth2.credentials import Credentials
    from google.auth.transport.requests import Request
    creds = Credentials(
        token=user.access_token, refresh_token=user.refresh_token,
        token_uri='https://oauth2.googleapis.com/token',
        client_id=app.config['GOOGLE_CLIENT_ID'],
        client_secret=app.config['GOOGLE_CLIENT_SECRET']
    )
    if creds.expired and creds.refresh_token:
        creds.refresh(Request())
        user.access_token = creds.token
        db.session.commit()
    return build('drive', 'v3', credentials=creds)


def get_or_create_folder(drive, name, parent_id=None):
    safe_name = (name or '').replace("\\", "\\\\").replace("'", "\\'")
    query = f"name='{safe_name}' and mimeType='application/vnd.google-apps.folder' and trashed=false"
    if parent_id:
        query += f" and '{parent_id}' in parents"
    files = drive.files().list(q=query, fields='files(id)', pageSize=1).execute().get('files', [])
    if files:
        return files[0]['id']
    meta = {'name': name, 'mimeType': 'application/vnd.google-apps.folder'}
    if parent_id:
        meta['parents'] = [parent_id]
    return drive.files().create(body=meta, fields='id').execute().get('id')


def drive_item_exists(drive, file_id):
    if not file_id:
        return False
    try:
        drive.files().get(fileId=file_id, fields='id').execute()
        return True
    except Exception:
        return False


def drive_safe_name(name, fallback='untitled'):
    clean = re.sub(r'[\\/:*?"<>|]+', ' ', (name or '').strip())
    clean = re.sub(r'\s+', ' ', clean).strip().rstrip('.')
    return clean[:120] or fallback


def find_drive_file_id(drive, name, parent_id):
    safe_name = (name or '').replace("\\", "\\\\").replace("'", "\\'")
    query = f"name='{safe_name}' and '{parent_id}' in parents and trashed=false"
    files = drive.files().list(q=query, fields='files(id)', pageSize=1).execute().get('files', [])
    return files[0]['id'] if files else ''


def get_valid_relationships(project_id):
    rels = CharacterRelationship.query.filter_by(project_id=project_id).all()
    invalid_rels = [
        rel for rel in rels
        if not rel.char_a or not rel.char_b
        or rel.char_a.project_id != project_id
        or rel.char_b.project_id != project_id
    ]
    if invalid_rels:
        for rel in invalid_rels:
            db.session.delete(rel)
        db.session.commit()
    invalid_ids = {rel.id for rel in invalid_rels}
    return [rel for rel in rels if rel.id not in invalid_ids]


def get_valid_lore_relationships(project_id):
    rels = LoreRelationship.query.filter_by(project_id=project_id).all()
    invalid_rels = [
        rel for rel in rels
        if not rel.lore_a or not rel.lore_b
        or rel.lore_a.project_id != project_id
        or rel.lore_b.project_id != project_id
    ]
    if invalid_rels:
        for rel in invalid_rels:
            db.session.delete(rel)
        db.session.commit()
    invalid_ids = {rel.id for rel in invalid_rels}
    return [rel for rel in rels if rel.id not in invalid_ids]


def create_drive_file(drive, name, content, parent_id):
    media  = MediaInMemoryUpload(content.encode('utf-8'), mimetype='text/plain', resumable=False)
    meta   = {'name': f"{drive_safe_name(name)}.txt", 'parents': [parent_id]}
    return drive.files().create(body=meta, media_body=media, fields='id').execute().get('id')


def update_drive_file(drive, file_id, name, content):
    media = MediaInMemoryUpload(content.encode('utf-8'), mimetype='text/plain', resumable=False)
    drive.files().update(fileId=file_id, body={'name': f"{drive_safe_name(name)}.txt"}, media_body=media).execute()


def ensure_drive_file_parent(drive, file_id, parent_id):
    if not file_id or not parent_id:
        return
    if not drive_item_exists(drive, file_id):
        return
    meta = drive.files().get(fileId=file_id, fields='parents').execute()
    parents = meta.get('parents', []) or []
    if parents == [parent_id]:
        return
    remove_parents = ','.join(parent for parent in parents if parent != parent_id)
    drive.files().update(
        fileId=file_id,
        addParents=parent_id if parent_id not in parents else None,
        removeParents=remove_parents or None,
        fields='id, parents'
    ).execute()


def upsert_drive_text_file(drive, name, content, parent_id, existing_file_id=''):
    file_name = f"{drive_safe_name(name)}.txt"
    file_id = existing_file_id if drive_item_exists(drive, existing_file_id) else ''
    if not file_id:
        file_id = find_drive_file_id(drive, file_name, parent_id)
    if file_id:
        try:
            update_drive_file(drive, file_id, name, content)
        except Exception:
            file_id = ''
    if file_id:
        ensure_drive_file_parent(drive, file_id, parent_id)
        return file_id
    return create_drive_file(drive, name, content, parent_id)


def delete_drive_file_if_exists(drive, file_id):
    if not file_id:
        return
    try:
        drive.files().delete(fileId=file_id).execute()
    except Exception as e:
        print(f"Drive delete error: {e}")


def fetch_image_payload(image_url):
    if not image_url:
        return None, None
    try:
        if image_url.startswith('data:'):
            header, encoded = image_url.split(',', 1)
            mime = header.split(';', 1)[0].split(':', 1)[1] if ':' in header else 'image/png'
            return base64.b64decode(encoded), mime
        response = http_requests.get(image_url, timeout=15)
        response.raise_for_status()
        return response.content, response.headers.get('Content-Type', 'image/png').split(';', 1)[0]
    except Exception as e:
        print(f"Image fetch error: {e}")
        return None, None


def image_extension_for_mime(mime_type):
    mapping = {
        'image/jpeg': '.jpg',
        'image/jpg': '.jpg',
        'image/png': '.png',
        'image/webp': '.webp',
        'image/gif': '.gif'
    }
    return mapping.get((mime_type or '').lower(), '.png')


def upsert_drive_image_file(drive, name, image_url, parent_id, existing_file_id=''):
    if not image_url:
        if existing_file_id:
            delete_drive_file_if_exists(drive, existing_file_id)
        return ''
    image_bytes, mime_type = fetch_image_payload(image_url)
    if not image_bytes:
        return existing_file_id or ''
    ext = image_extension_for_mime(mime_type)
    file_name = f"{drive_safe_name(name)} image{ext}"
    media = MediaInMemoryUpload(image_bytes, mimetype=mime_type or 'image/png', resumable=False)
    file_id = existing_file_id if drive_item_exists(drive, existing_file_id) else ''
    if not file_id:
        file_id = find_drive_file_id(drive, file_name, parent_id)
    if file_id:
        try:
            drive.files().update(fileId=file_id, body={'name': file_name}, media_body=media).execute()
        except Exception:
            file_id = ''
    if file_id:
        ensure_drive_file_parent(drive, file_id, parent_id)
        return file_id
    return drive.files().create(
        body={'name': file_name, 'parents': [parent_id]},
        media_body=media,
        fields='id'
    ).execute().get('id')


def html_to_plain_text(html):
    text = re.sub(r'<br\s*/?>', '\n', html)
    text = re.sub(r'</p>|</h[1-3]>', '\n\n', text)
    text = re.sub(r'<li>', '• ', text)
    text = re.sub(r'</li>', '\n', text)
    text = re.sub(r'<[^>]+>', '', text)
    return re.sub(r'\n{3,}', '\n\n', text).strip()


def setup_scripvia_folder(user):
    if user.scripvia_folder_id:
        try:
            drive = get_drive_service(user)
            if drive_item_exists(drive, user.scripvia_folder_id):
                return user.scripvia_folder_id
            user.scripvia_folder_id = ''
            db.session.commit()
        except Exception:
            pass
    if user.scripvia_folder_id:
        return user.scripvia_folder_id
    try:
        drive = get_drive_service(user)
        folder_id = get_or_create_folder(drive, 'Scripvia')
        user.scripvia_folder_id = folder_id
        db.session.commit()
        return folder_id
    except Exception as e:
        print(f"Could not create Scripvia folder: {e}")
        return None


def ensure_project_drive_structure(drive, user, project):
    if user.scripvia_folder_id and not drive_item_exists(drive, user.scripvia_folder_id):
        user.scripvia_folder_id = ''
        db.session.commit()
    if project.drive_folder_id and not drive_item_exists(drive, project.drive_folder_id):
        project.drive_folder_id = ''
        db.session.commit()
    if not user.scripvia_folder_id:
        user.scripvia_folder_id = get_or_create_folder(drive, 'Scripvia')
        db.session.commit()
    if not project.drive_folder_id:
        project.drive_folder_id = get_or_create_folder(drive, drive_safe_name(project.title), parent_id=user.scripvia_folder_id)
        db.session.commit()

    folders = {
        'project': project.drive_folder_id,
        'chapters': get_or_create_folder(drive, 'Chapters', parent_id=project.drive_folder_id)
    }
    if project.genre in CREATIVE_GENRES:
        folders['characters'] = get_or_create_folder(drive, 'Characters', parent_id=project.drive_folder_id)
        folders['scenes'] = get_or_create_folder(drive, 'Scenes', parent_id=project.drive_folder_id)
        folders['lore'] = get_or_create_folder(drive, 'Lore', parent_id=project.drive_folder_id)
        folders['web'] = get_or_create_folder(drive, 'Web', parent_id=project.drive_folder_id)
    return folders


def build_document_drive_content(doc):
    return f"{doc.title}\n{'=' * len(doc.title)}\n\n{html_to_plain_text(doc.content or '')}"


def append_drive_metadata(content, payload):
    return (
        f"{content.rstrip()}\n\n"
        f"--- SCRIPVIA JSON ---\n"
        f"{json.dumps(payload, ensure_ascii=False)}\n"
        f"--- END SCRIPVIA JSON ---"
    )


def extract_drive_metadata(content):
    marker = "\n--- SCRIPVIA JSON ---\n"
    end_marker = "\n--- END SCRIPVIA JSON ---"
    if marker not in content or end_marker not in content:
        return content, {}
    prefix, suffix = content.split(marker, 1)
    json_blob, _ = suffix.split(end_marker, 1)
    try:
        return prefix.rstrip(), json.loads(json_blob.strip())
    except Exception:
        return content, {}


def build_character_drive_content(char):
    titles = ', '.join(char.get_titles()) or 'None'
    body = (
        f"Name: {char.name}\n"
        f"Titles: {titles}\n"
        f"Role: {char.role or 'None'}\n"
        f"Age: {char.age or 'Unknown'}\n\n"
        f"Appearance:\n{char.appearance or ''}\n\n"
        f"Personality:\n{char.personality or ''}\n\n"
        f"Backstory:\n{char.backstory or ''}\n\n"
        f"Notes:\n{char.extra_notes or ''}"
    )
    return append_drive_metadata(body, {
        'type': 'character',
        'name': char.name,
        'role': char.role,
        'age': char.age,
        'appearance': char.appearance,
        'personality': char.personality,
        'backstory': char.backstory,
        'extra_notes': char.extra_notes,
        'titles': char.get_titles(),
        'image_url': char.image_url or '',
        'image_focus': char.image_focus or 'center',
        'web_x': char.web_x,
        'web_y': char.web_y
    })


def build_scene_drive_content(scene):
    return (
        f"Scene: {scene.title}\n"
        f"Mood: {scene.mood or 'None'}\n"
        f"Connected Chapter: {scene.connected_chapter or 'None'}\n\n"
        f"{html_to_plain_text(scene.content or '')}"
    )


def build_lore_drive_content(item):
    body = (
        f"Name: {item.name}\n"
        f"Category: {item.category or 'item'}\n\n"
        f"Description:\n{item.description or ''}\n\n"
        f"Timeline Date: {item.event_date or 'None'}\n"
        f"Timeline Order: {item.event_order or 0}\n"
        f"Map Position: ({item.map_x:.1f}, {item.map_y:.1f})\n\n"
        f"Notes:\n{item.extra_notes or ''}"
    )
    return append_drive_metadata(body, {
        'type': 'lore',
        'name': item.name,
        'category': item.category,
        'description': item.description,
        'aliases': item.get_aliases(),
        'image_url': item.image_url or '',
        'image_focus': item.image_focus or 'center',
        'extra_notes': item.extra_notes,
        'event_date': item.event_date or '',
        'event_order': item.event_order or 0,
        'show_in_web': bool(item.show_in_web),
        'web_x': item.web_x,
        'web_y': item.web_y,
        'map_x': item.map_x,
        'map_y': item.map_y
    })


def build_relationships_drive_content(project_id):
    rels = get_valid_relationships(project_id)
    if not rels:
        return "No relationships yet."
    return '\n'.join(
        f"{rel.char_a.name if rel.char_a else ''} {format_relation_text(rel.relation_type)} {rel.char_b.name if rel.char_b else ''}".strip()
        for rel in rels
    )


def build_lore_relationships_drive_content(project_id):
    rels = get_valid_lore_relationships(project_id)
    if not rels:
        return "No lore connections yet."
    return '\n'.join(
        f"{rel.lore_a.name if rel.lore_a else ''} {format_relation_text(rel.relation_type)} {rel.lore_b.name if rel.lore_b else ''}".strip()
        for rel in rels
    )


def format_relation_text(value):
    return (value or 'is connected to').strip()


def sync_project_relationships_to_drive(user, project):
    drive = get_drive_service(user)
    folders = ensure_project_drive_structure(drive, user, project)
    if 'web' not in folders:
        return
    upsert_drive_text_file(
        drive,
        'relationships',
        build_relationships_drive_content(project.id),
        folders['web']
    )


def sync_project_lore_relationships_to_drive(user, project):
    drive = get_drive_service(user)
    folders = ensure_project_drive_structure(drive, user, project)
    if 'web' not in folders:
        return
    upsert_drive_text_file(
        drive,
        'lore_connections',
        build_lore_relationships_drive_content(project.id),
        folders['web']
    )


def sync_full_project_to_drive(user, project):
    drive = get_drive_service(user)
    folders = ensure_project_drive_structure(drive, user, project)

    for doc in Document.query.filter_by(project_id=project.id).order_by(Document.order_index.asc(), Document.created_at.asc()).all():
        doc.drive_file_id = upsert_drive_text_file(
            drive, doc.title, build_document_drive_content(doc), folders['chapters'], doc.drive_file_id
        )

    if project.genre in CREATIVE_GENRES:
        characters_folder = folders.get('characters')
        scenes_folder = folders.get('scenes')
        lore_folder = folders.get('lore')

        for char in Character.query.filter_by(project_id=project.id).order_by(Character.name.asc()).all():
            if characters_folder:
                char.drive_file_id = upsert_drive_text_file(
                    drive, char.name, build_character_drive_content(char), characters_folder, char.drive_file_id
                )
                char.drive_image_file_id = upsert_drive_image_file(
                    drive, char.name, char.image_url, characters_folder, char.drive_image_file_id
                )

        for scene in Scene.query.filter_by(project_id=project.id).order_by(Scene.updated_at.desc()).all():
            if scenes_folder:
                scene.drive_file_id = upsert_drive_text_file(
                    drive, scene.title, build_scene_drive_content(scene), scenes_folder, scene.drive_file_id
                )

        for item in LoreItem.query.filter_by(project_id=project.id).order_by(LoreItem.name.asc()).all():
            if lore_folder:
                item.drive_file_id = upsert_drive_text_file(
                    drive, item.name, build_lore_drive_content(item), lore_folder, item.drive_file_id
                )
                item.drive_image_file_id = upsert_drive_image_file(
                    drive, item.name, item.image_url, lore_folder, item.drive_image_file_id
                )

        sync_project_relationships_to_drive(user, project)

    db.session.commit()


def list_drive_children(drive, parent_id, mime_type=None):
    query = f"'{parent_id}' in parents and trashed=false"
    if mime_type:
        query += f" and mimeType='{mime_type}'"
    page_token = None
    items = []
    while True:
        response = drive.files().list(
            q=query,
            fields='nextPageToken, files(id, name, mimeType)',
            pageSize=200,
            pageToken=page_token,
            orderBy='name'
        ).execute()
        items.extend(response.get('files', []))
        page_token = response.get('nextPageToken')
        if not page_token:
            break
    return items


def download_drive_file_text(drive, file_id):
    payload = drive.files().get_media(fileId=file_id).execute()
    if isinstance(payload, bytes):
        return payload.decode('utf-8', errors='replace')
    return str(payload or '')


def download_drive_file_as_data_url(drive, file_id, fallback_name='image.png'):
    meta = drive.files().get(fileId=file_id, fields='name, mimeType').execute()
    raw = drive.files().get_media(fileId=file_id).execute()
    mime = meta.get('mimeType') or 'image/png'
    encoded = base64.b64encode(raw).decode('ascii')
    return f"data:{mime};base64,{encoded}"


def extract_named_section(content, heading):
    pattern = rf"{re.escape(heading)}:\n(.*?)(?:\n[A-Z][A-Za-z ]+:\n|\Z)"
    match = re.search(pattern, content, flags=re.S)
    return match.group(1).strip() if match else ''


def parse_legacy_character_content(name, content):
    lines = [line.rstrip() for line in content.splitlines()]
    title_line = next((line for line in lines if line.startswith('Titles:')), '')
    role_line = next((line for line in lines if line.startswith('Role:')), '')
    age_line = next((line for line in lines if line.startswith('Age:')), '')
    titles = [part.strip() for part in title_line.replace('Titles:', '', 1).split(',') if part.strip() and part.strip().lower() != 'none']
    return {
        'name': name,
        'titles': titles,
        'role': role_line.replace('Role:', '', 1).strip().replace('None', ''),
        'age': age_line.replace('Age:', '', 1).strip().replace('Unknown', ''),
        'appearance': extract_named_section(content, 'Appearance'),
        'personality': extract_named_section(content, 'Personality'),
        'backstory': extract_named_section(content, 'Backstory'),
        'extra_notes': extract_named_section(content, 'Notes'),
        'image_url': '',
        'image_focus': 'center'
    }


def parse_legacy_lore_content(name, content):
    lines = [line.rstrip() for line in content.splitlines()]
    category_line = next((line for line in lines if line.startswith('Category:')), '')
    date_line = next((line for line in lines if line.startswith('Timeline Date:')), '')
    order_line = next((line for line in lines if line.startswith('Timeline Order:')), '')
    map_line = next((line for line in lines if line.startswith('Map Position:')), '')
    coords = re.search(r'\(([-\d.]+),\s*([-\d.]+)\)', map_line or '')
    return {
        'name': name,
        'category': category_line.replace('Category:', '', 1).strip() or 'item',
        'description': extract_named_section(content, 'Description'),
        'aliases': [],
        'image_url': '',
        'image_focus': 'center',
        'extra_notes': extract_named_section(content, 'Notes'),
        'event_date': date_line.replace('Timeline Date:', '', 1).strip().replace('None', ''),
        'event_order': int((order_line.replace('Timeline Order:', '', 1).strip() or '0')),
        'show_in_web': True,
        'web_x': None,
        'web_y': None,
        'map_x': float(coords.group(1)) if coords else 50.0,
        'map_y': float(coords.group(2)) if coords else 50.0
    }


def plain_text_to_html(text):
    paragraphs = [segment.strip() for segment in re.split(r'\n{2,}', text.strip()) if segment.strip()]
    if not paragraphs:
        return ''
    return ''.join(f"<p>{segment.replace(chr(10), '<br>')}</p>" for segment in paragraphs)


def import_drive_project(drive, user, folder_meta):
    project = Project.query.filter_by(user_id=user.id, drive_folder_id=folder_meta['id']).first()
    if not project:
        project = Project(
            title=folder_meta['name'],
            description='Imported from Google Drive',
            genre='fantasy',
            user_id=user.id,
            guest_session_id='',
            drive_folder_id=folder_meta['id']
        )
        db.session.add(project)
        db.session.commit()

    child_folders = {item['name'].lower(): item for item in list_drive_children(drive, folder_meta['id'], 'application/vnd.google-apps.folder')}
    is_creative_project = any(key in child_folders for key in ['characters', 'scenes', 'lore'])
    target_genre = project.genre if project.genre and project.genre != 'general' else ('fantasy' if is_creative_project else 'general')
    if project.genre != target_genre:
        project.genre = target_genre
        db.session.commit()

    chapters_folder = child_folders.get('chapters')
    if chapters_folder:
        for file_meta in list_drive_children(drive, chapters_folder['id']):
            if file_meta.get('mimeType') == 'application/vnd.google-apps.folder' or not file_meta['name'].lower().endswith('.txt'):
                continue
            existing = Document.query.filter_by(project_id=project.id, drive_file_id=file_meta['id']).first()
            if existing:
                continue
            raw = download_drive_file_text(drive, file_meta['id'])
            title = re.sub(r'\.txt$', '', file_meta['name'], flags=re.I)
            content_text = raw
            if '\n\n' in raw:
                _, content_text = raw.split('\n\n', 1)
            doc = Document(
                title=title,
                content=plain_text_to_html(content_text),
                project_id=project.id,
                drive_file_id=file_meta['id'],
                order_index=Document.query.filter_by(project_id=project.id).count()
            )
            db.session.add(doc)

    characters_folder = child_folders.get('characters')
    if characters_folder:
        image_files = {item['name'].rsplit(' image', 1)[0].strip().lower(): item for item in list_drive_children(drive, characters_folder['id']) if 'image/' in (item.get('mimeType') or '')}
        for file_meta in list_drive_children(drive, characters_folder['id']):
            if file_meta.get('mimeType') == 'application/vnd.google-apps.folder' or not file_meta['name'].lower().endswith('.txt'):
                continue
            existing = Character.query.filter_by(project_id=project.id, drive_file_id=file_meta['id']).first()
            if existing:
                continue
            raw = download_drive_file_text(drive, file_meta['id'])
            _, metadata = extract_drive_metadata(raw)
            name = re.sub(r'\.txt$', '', file_meta['name'], flags=re.I)
            payload = metadata if metadata.get('type') == 'character' else parse_legacy_character_content(name, raw)
            image_meta = image_files.get((payload.get('name') or name).strip().lower())
            image_url = payload.get('image_url') or (download_drive_file_as_data_url(drive, image_meta['id']) if image_meta else '')
            char = Character(
                project_id=project.id,
                name=payload.get('name') or name,
                role=payload.get('role', ''),
                age=payload.get('age', ''),
                personality=payload.get('personality', ''),
                backstory=payload.get('backstory', ''),
                appearance=payload.get('appearance', ''),
                image_url=image_url,
                image_focus=payload.get('image_focus', 'center'),
                extra_notes=payload.get('extra_notes', ''),
                drive_file_id=file_meta['id'],
                drive_image_file_id=image_meta['id'] if image_meta else '',
                web_x=payload.get('web_x'),
                web_y=payload.get('web_y')
            )
            char.set_titles(payload.get('titles', []))
            db.session.add(char)

    scenes_folder = child_folders.get('scenes')
    if scenes_folder:
        for file_meta in list_drive_children(drive, scenes_folder['id']):
            if file_meta.get('mimeType') == 'application/vnd.google-apps.folder' or not file_meta['name'].lower().endswith('.txt'):
                continue
            existing = Scene.query.filter_by(project_id=project.id, drive_file_id=file_meta['id']).first()
            if existing:
                continue
            raw = download_drive_file_text(drive, file_meta['id'])
            scene = Scene(
                project_id=project.id,
                title=re.sub(r'\.txt$', '', file_meta['name'], flags=re.I),
                content=plain_text_to_html(raw.split('\n\n', 1)[1] if '\n\n' in raw else raw),
                mood='',
                connected_chapter='',
                drive_file_id=file_meta['id']
            )
            db.session.add(scene)

    lore_folder = child_folders.get('lore')
    if lore_folder:
        image_files = {item['name'].rsplit(' image', 1)[0].strip().lower(): item for item in list_drive_children(drive, lore_folder['id']) if 'image/' in (item.get('mimeType') or '')}
        for file_meta in list_drive_children(drive, lore_folder['id']):
            if file_meta.get('mimeType') == 'application/vnd.google-apps.folder' or not file_meta['name'].lower().endswith('.txt'):
                continue
            existing = LoreItem.query.filter_by(project_id=project.id, drive_file_id=file_meta['id']).first()
            if existing:
                continue
            raw = download_drive_file_text(drive, file_meta['id'])
            _, metadata = extract_drive_metadata(raw)
            name = re.sub(r'\.txt$', '', file_meta['name'], flags=re.I)
            payload = metadata if metadata.get('type') == 'lore' else parse_legacy_lore_content(name, raw)
            image_meta = image_files.get((payload.get('name') or name).strip().lower())
            image_url = payload.get('image_url') or (download_drive_file_as_data_url(drive, image_meta['id']) if image_meta else '')
            lore = LoreItem(
                project_id=project.id,
                name=payload.get('name') or name,
                category=payload.get('category', 'item'),
                description=payload.get('description', ''),
                image_url=image_url,
                image_focus=payload.get('image_focus', 'center'),
                extra_notes=payload.get('extra_notes', ''),
                event_date=payload.get('event_date', ''),
                event_order=int(payload.get('event_order') or 0),
                show_in_web=payload.get('show_in_web', True),
                web_x=payload.get('web_x'),
                web_y=payload.get('web_y'),
                map_x=float(payload.get('map_x') or 50),
                map_y=float(payload.get('map_y') or 50),
                drive_file_id=file_meta['id'],
                drive_image_file_id=image_meta['id'] if image_meta else ''
            )
            lore.set_aliases(payload.get('aliases', []))
            db.session.add(lore)

    db.session.commit()
    return project


def import_projects_from_drive(user):
    if not user or not user.access_token:
        return
    drive = get_drive_service(user)
    root_folder_id = setup_scripvia_folder(user)
    if not root_folder_id:
        return
    project_folders = list_drive_children(drive, root_folder_id, 'application/vnd.google-apps.folder')
    for folder in project_folders:
        import_drive_project(drive, user, folder)


def drive_bg(fn, *args):
    """Run a Drive operation in a background thread."""
    t = threading.Thread(target=fn, args=args)
    t.daemon = True
    t.start()


# --- ROUTES ---

@app.route('/')
def index():
    # Try to restore session from file if flask session is empty
    if not session.get('user_id') and not session.get('guest_session_id'):
        saved = load_persistent_session()
        if saved:
            if saved.get('user_id'):
                session['user_id'] = saved['user_id']
                session.permanent = True
            elif saved.get('is_guest') and saved.get('guest_session_id'):
                session['guest_session_id'] = saved['guest_session_id']
                session.permanent = True
        else:
            # Nothing saved at all → show login
            return redirect('/login')
    return render_template('index.html')

@app.route('/login')
def login_page():
    return redirect('/') if get_current_user() else render_template('login.html')


# --- AUTH ---

@app.route('/auth/login')
def auth_login():
    state = secrets.token_urlsafe(32)
    session['oauth_state'] = state
    params = {
        'client_id': app.config['GOOGLE_CLIENT_ID'],
        'redirect_uri': app.config['GOOGLE_REDIRECT_URI'],
        'response_type': 'code', 'scope': ' '.join(app.config['GOOGLE_SCOPES']),
        'access_type': 'offline', 'prompt': 'consent', 'state': state
    }
    url = 'https://accounts.google.com/o/oauth2/auth?' + urllib.parse.urlencode(params)
    print('[DEBUG AUTH URL]', url)  # ← add this
    return redirect(url)

@app.route('/auth/callback')
def auth_callback():
    code = request.args.get('code')
    if not code:
        return 'Login failed', 400
    tokens = http_requests.post('https://oauth2.googleapis.com/token', data={
        'code': code, 'client_id': app.config['GOOGLE_CLIENT_ID'],
        'client_secret': app.config['GOOGLE_CLIENT_SECRET'],
        'redirect_uri': app.config['GOOGLE_REDIRECT_URI'], 'grant_type': 'authorization_code'
    }).json()
    if 'error' in tokens:
        return f"Token error: {tokens.get('error_description', tokens['error'])}", 400
    profile = http_requests.get(
        'https://www.googleapis.com/oauth2/v2/userinfo',
        headers={'Authorization': f"Bearer {tokens['access_token']}"}
    ).json()
    user = User.query.filter_by(google_id=profile['id']).first()
    if not user:
        user = User(google_id=profile['id'], email=profile.get('email', ''),
                    name=profile.get('name', ''), picture=profile.get('picture', ''))
        db.session.add(user)
    user.access_token  = tokens.get('access_token', '')
    user.refresh_token = tokens.get('refresh_token', getattr(user, 'refresh_token', ''))
    db.session.commit()
    guest_session_id = session.get('guest_session_id')
    if guest_session_id:
        guest_projects = Project.query.filter_by(user_id=None, guest_session_id=guest_session_id).all()
        for project in guest_projects:
            project.user_id = user.id
            project.guest_session_id = ''
        if guest_projects:
            db.session.commit()
    session['user_id'] = user.id
    save_persistent_session(user_id=user.id)
    session.permanent  = True
    save_persistent_session(user.id)
    setup_scripvia_folder(user)
    return redirect('/?logged_in=true')


@app.route('/auth/logout')
def auth_logout():
    session.clear()
    clear_persistent_session()
    return redirect('/login')


@app.route('/auth/me')
def auth_me():
    user = get_current_user()
    if user:
        return jsonify({'logged_in': True, 'user': user.to_dict(), 'scope_key': f"user_{user.id}"})
    return jsonify({'logged_in': False, 'scope_key': f"guest_{get_guest_session_id()}"})

@app.route('/auth/guest', methods=['POST'])
def auth_guest():
    data = request.get_json()
    name = (data or {}).get('name', 'Guest').strip() or 'Guest'
    gid  = get_guest_session_id()  # creates or restores guest session
    session['guest_name'] = name
    session.permanent = True
    save_persistent_session(is_guest=True, guest_session_id=gid)
    return jsonify({'ok': True})

# --- PROJECTS ---

@app.route('/api/projects', methods=['GET'])
def get_projects():
    user = get_current_user()
    if user:
        should_attempt_import = (
            bool(user.access_token) and (
                session.get('drive_import_checked_user_id') != user.id or
                Project.query.filter_by(user_id=user.id).count() == 0
            )
        )
        if should_attempt_import:
            try:
                import_projects_from_drive(user)
                session['drive_import_checked_user_id'] = user.id
                session.permanent = True
            except Exception as e:
                print(f"Drive import error: {e}")
        projects = Project.query.filter(
            (Project.user_id == user.id) | legacy_project_filter()
        ).order_by(Project.updated_at.desc()).all()
    else:
        projects = Project.query.filter(
            ((Project.user_id.is_(None)) & (Project.guest_session_id == get_guest_session_id())) | legacy_project_filter()
        ).order_by(Project.updated_at.desc()).all()
    return jsonify([p.to_dict() for p in projects])


@app.route('/api/projects/<int:project_id>/stats', methods=['GET'])
def get_project_stats(project_id):
    project     = get_project_or_404(project_id)
    total_words = sum(len(html_to_plain_text(d.content or '').split()) for d in project.documents if d.content)
    scene_words = sum(len(html_to_plain_text(s.content or '').split()) for s in project.scenes if s.content)
    all_dates   = [project.updated_at] + [d.updated_at for d in project.documents]
    return jsonify({
        'id': project.id, 'title': project.title, 'description': project.description,
        'genre': project.genre, 'is_creative': project.genre in CREATIVE_GENRES,
        'chapter_count': len(project.documents), 'character_count': len(project.characters),
        'scene_count': len(project.scenes), 'lore_count': len(project.lore_items),
        'total_words': total_words + scene_words,
        'last_edited': max(all_dates).isoformat() if all_dates else None,
        'created_at': project.created_at.isoformat()
    })


@app.route('/api/projects', methods=['POST'])
def create_project():
    data = request.get_json()
    if not data or not data.get('title'):
        return jsonify({'error': 'Title required'}), 400
    user = get_current_user()
    project = Project(
        title=data['title'],
        description=data.get('description', ''),
        genre=data.get('genre', 'general'),
        map_image_url=data.get('map_image_url', ''),
        user_id=user.id if user else None,
        guest_session_id='' if user else get_guest_session_id()
    )
    db.session.add(project)
    db.session.commit()

    if user and user.access_token:
        def _bg(app, pid, uid):
            with app.app_context():
                u, p = User.query.get(uid), Project.query.get(pid)
                if not u or not p: return
                try:
                    drive = get_drive_service(u)
                    ensure_project_drive_structure(drive, u, p)
                except Exception as e:
                    print(f"Drive folder error: {e}")
        drive_bg(_bg, app, project.id, user.id)

    return jsonify(project.to_dict()), 201


@app.route('/api/projects/<int:project_id>', methods=['PUT'])
def update_project(project_id):
    project = get_project_or_404(project_id)
    data = request.get_json() or {}
    for field in ['title', 'description', 'genre', 'map_image_url']:
        if field in data:
            setattr(project, field, data[field])
    db.session.commit()
    return jsonify(project.to_dict())


@app.route('/api/projects/<int:project_id>', methods=['DELETE'])
def delete_project(project_id):
    project = get_project_or_404(project_id)
    folder_id = project.drive_folder_id
    db.session.delete(project)
    db.session.commit()

    user = get_current_user()
    if user and user.access_token and folder_id:
        def _bg(app, fid, uid):
            with app.app_context():
                u = User.query.get(uid)
                if not u: return
                try:
                    get_drive_service(u).files().delete(fileId=fid).execute()
                except Exception as e:
                    print(f"Drive folder delete error: {e}")
        drive_bg(_bg, app, folder_id, user.id)

    return jsonify({'message': 'Deleted'})


# --- DOCUMENTS ---

@app.route('/api/projects/<int:project_id>/documents', methods=['GET'])
def get_documents(project_id):
    get_project_or_404(project_id)
    docs = Document.query.filter_by(project_id=project_id).order_by(
        Document.order_index.asc(), Document.created_at.asc()).all()
    return jsonify([d.to_dict() for d in docs])


@app.route('/api/projects/<int:project_id>/documents/reorder', methods=['POST'])
def reorder_documents(project_id):
    get_project_or_404(project_id)
    data = request.get_json()
    if not data or 'order' not in data:
        return jsonify({'error': 'order required'}), 400
    for i, doc_id in enumerate(data['order']):
        doc = Document.query.get(doc_id)
        if doc and doc.project_id == project_id:
            doc.order_index = i
    db.session.commit()
    return jsonify({'message': 'Reordered'})


@app.route('/api/projects/<int:project_id>/documents', methods=['POST'])
def create_document(project_id):
    project = get_project_or_404(project_id)
    data    = request.get_json()
    if not data or not data.get('title'):
        return jsonify({'error': 'Title required'}), 400
    last  = Document.query.filter_by(project_id=project_id).order_by(Document.order_index.desc()).first()
    doc   = Document(title=data['title'], content=data.get('content', ''),
                     project_id=project_id, drive_file_id='',
                     order_index=(last.order_index + 1) if last else 0)
    db.session.add(doc)
    db.session.commit()

    user = get_current_user()
    if user and user.access_token:
        def _bg(app, did, pid, uid):
            with app.app_context():
                u, p, d = User.query.get(uid), Project.query.get(pid), Document.query.get(did)
                if not u or not p or not d: return
                try:
                    drive = get_drive_service(u)
                    folders = ensure_project_drive_structure(drive, u, p)
                    d.drive_file_id = upsert_drive_text_file(
                        drive, d.title, build_document_drive_content(d), folders['chapters'], d.drive_file_id
                    )
                    db.session.commit()
                except Exception as e:
                    print(f"Drive file error: {e}")
        drive_bg(_bg, app, doc.id, project.id, user.id)

    return jsonify(doc.to_dict()), 201


@app.route('/api/documents/<int:doc_id>', methods=['GET'])
def get_document(doc_id):
    doc = Document.query.get_or_404(doc_id)
    get_project_or_404(doc.project_id)
    return jsonify(doc.to_dict())


@app.route('/api/documents/<int:doc_id>', methods=['PUT'])
def update_document(doc_id):
    doc  = Document.query.get_or_404(doc_id)
    get_project_or_404(doc.project_id)
    data = request.get_json()
    if 'title'   in data: doc.title   = data['title']
    if 'content' in data: doc.content = data['content']
    doc.updated_at = datetime.utcnow()
    db.session.commit()
    return jsonify(doc.to_dict())


@app.route('/api/documents/<int:doc_id>', methods=['DELETE'])
def delete_document(doc_id):
    doc = Document.query.get_or_404(doc_id)
    get_project_or_404(doc.project_id)
    file_id = doc.drive_file_id
    db.session.delete(doc)
    db.session.commit()

    user = get_current_user()
    if user and user.access_token and file_id:
        def _bg(app, fid, uid):
            with app.app_context():
                u = User.query.get(uid)
                if not u: return
                delete_drive_file_if_exists(get_drive_service(u), fid)
        drive_bg(_bg, app, file_id, user.id)

    return jsonify({'message': 'Deleted'})


# --- CHARACTERS ---

@app.route('/api/projects/<int:project_id>/characters', methods=['GET'])
def get_characters(project_id):
    get_project_or_404(project_id)
    return jsonify([c.to_dict() for c in Character.query.filter_by(project_id=project_id).order_by(Character.name).all()])


@app.route('/api/projects/<int:project_id>/characters', methods=['POST'])
def create_character(project_id):
    get_project_or_404(project_id)
    data = request.get_json()
    if not data or not data.get('name'):
        return jsonify({'error': 'Name required'}), 400
    char = Character(project_id=project_id, name=data['name'], role=data.get('role', ''),
                     age=data.get('age', ''), personality=data.get('personality', ''),
                     backstory=data.get('backstory', ''), appearance=data.get('appearance', ''),
                     image_url=data.get('image_url', ''), image_focus=data.get('image_focus', 'center'),
                     extra_notes=data.get('extra_notes', ''),
                     web_x=data.get('web_x'), web_y=data.get('web_y'))
    char.set_titles(data.get('titles', []))
    db.session.add(char)
    db.session.commit()

    user = get_current_user()
    if user and user.access_token:
        def _bg(app, cid, pid, uid):
            with app.app_context():
                u, p, c = User.query.get(uid), Project.query.get(pid), Character.query.get(cid)
                if not u or not p or not c: return
                try:
                    drive = get_drive_service(u)
                    folders = ensure_project_drive_structure(drive, u, p)
                    folder = folders.get('characters')
                    if not folder:
                        return
                    c.drive_file_id = upsert_drive_text_file(
                        drive, c.name, build_character_drive_content(c), folder, c.drive_file_id
                    )
                    c.drive_image_file_id = upsert_drive_image_file(
                        drive, c.name, c.image_url, folder, c.drive_image_file_id
                    )
                    sync_project_relationships_to_drive(u, p)
                    db.session.commit()
                except Exception as e:
                    print(f"Character Drive sync error: {e}")
        drive_bg(_bg, app, char.id, project_id, user.id)

    return jsonify(char.to_dict()), 201


@app.route('/api/characters/<int:char_id>', methods=['GET'])
def get_character(char_id):
    char = Character.query.get_or_404(char_id)
    get_project_or_404(char.project_id)
    return jsonify(char.to_dict())


@app.route('/api/characters/<int:char_id>', methods=['PUT'])
def update_character(char_id):
    char = Character.query.get_or_404(char_id)
    get_project_or_404(char.project_id)
    data = request.get_json()
    for f in ['name','role','age','personality','backstory','appearance','image_url','image_focus','extra_notes','web_x','web_y']:
        if f in data: setattr(char, f, data[f])
    if 'titles' in data:
        char.set_titles(data.get('titles', []))
    db.session.commit()

    user = get_current_user()
    if user and user.access_token:
        def _bg(app, cid, uid):
            with app.app_context():
                u, c = User.query.get(uid), Character.query.get(cid)
                if not u or not c:
                    return
                p = Project.query.get(c.project_id)
                if not p:
                    return
                try:
                    drive = get_drive_service(u)
                    folders = ensure_project_drive_structure(drive, u, p)
                    folder = folders.get('characters')
                    if not folder:
                        return
                    c.drive_file_id = upsert_drive_text_file(
                        drive, c.name, build_character_drive_content(c), folder, c.drive_file_id
                    )
                    c.drive_image_file_id = upsert_drive_image_file(
                        drive, c.name, c.image_url, folder, c.drive_image_file_id
                    )
                    sync_project_relationships_to_drive(u, p)
                    db.session.commit()
                except Exception as e:
                    print(f"Character Drive update error: {e}")
        drive_bg(_bg, app, char.id, user.id)
    return jsonify(char.to_dict())


@app.route('/api/characters/<int:char_id>', methods=['DELETE'])
def delete_character(char_id):
    char = Character.query.get_or_404(char_id)
    get_project_or_404(char.project_id)
    drive_file_id = char.drive_file_id
    drive_image_file_id = char.drive_image_file_id
    project_id = char.project_id
    CharacterRelationship.query.filter(
        (CharacterRelationship.char_a_id == char.id) | (CharacterRelationship.char_b_id == char.id)
    ).delete(synchronize_session=False)
    db.session.delete(char)
    db.session.commit()

    user = get_current_user()
    if user and user.access_token:
        def _bg(app, uid, pid, text_fid, image_fid):
            with app.app_context():
                u = User.query.get(uid)
                p = Project.query.get(pid)
                if not u:
                    return
                drive = get_drive_service(u)
                delete_drive_file_if_exists(drive, text_fid)
                delete_drive_file_if_exists(drive, image_fid)
                if p:
                    sync_project_relationships_to_drive(u, p)
        drive_bg(_bg, app, user.id, project_id, drive_file_id, drive_image_file_id)
    return jsonify({'message': 'Deleted'})


# --- SCENES ---

@app.route('/api/projects/<int:project_id>/scenes', methods=['GET'])
def get_scenes(project_id):
    get_project_or_404(project_id)
    return jsonify([s.to_dict() for s in Scene.query.filter_by(project_id=project_id).order_by(Scene.updated_at.desc()).all()])


@app.route('/api/projects/<int:project_id>/scenes', methods=['POST'])
def create_scene(project_id):
    get_project_or_404(project_id)
    data = request.get_json()
    if not data or not data.get('title'):
        return jsonify({'error': 'Title required'}), 400
    scene = Scene(project_id=project_id, title=data['title'], content=data.get('content', ''),
                  mood=data.get('mood', ''), connected_chapter=data.get('connected_chapter', ''))
    db.session.add(scene)
    db.session.commit()

    user = get_current_user()
    if user and user.access_token:
        def _bg(app, sid, pid, uid):
            with app.app_context():
                u, p, s = User.query.get(uid), Project.query.get(pid), Scene.query.get(sid)
                if not u or not p or not s: return
                try:
                    drive = get_drive_service(u)
                    folders = ensure_project_drive_structure(drive, u, p)
                    folder = folders.get('scenes')
                    if not folder:
                        return
                    s.drive_file_id = upsert_drive_text_file(
                        drive, s.title, build_scene_drive_content(s), folder, s.drive_file_id
                    )
                    db.session.commit()
                except Exception as e:
                    print(f"Scene Drive sync error: {e}")
        drive_bg(_bg, app, scene.id, scene.project_id, user.id)

    return jsonify(scene.to_dict()), 201


@app.route('/api/scenes/<int:scene_id>', methods=['GET'])
def get_scene(scene_id):
    scene = Scene.query.get_or_404(scene_id)
    get_project_or_404(scene.project_id)
    return jsonify(scene.to_dict())


@app.route('/api/scenes/<int:scene_id>', methods=['PUT'])
def update_scene(scene_id):
    scene = Scene.query.get_or_404(scene_id)
    get_project_or_404(scene.project_id)
    data  = request.get_json()
    for f in ['title','content','mood','connected_chapter']:
        if f in data: setattr(scene, f, data[f])
    scene.updated_at = datetime.utcnow()
    db.session.commit()

    user = get_current_user()
    if user and user.access_token:
        def _bg(app, sid, uid):
            with app.app_context():
                u, s = User.query.get(uid), Scene.query.get(sid)
                if not u or not s:
                    return
                p = Project.query.get(s.project_id)
                if not p:
                    return
                try:
                    drive = get_drive_service(u)
                    folders = ensure_project_drive_structure(drive, u, p)
                    folder = folders.get('scenes')
                    if not folder:
                        return
                    s.drive_file_id = upsert_drive_text_file(
                        drive, s.title, build_scene_drive_content(s), folder, s.drive_file_id
                    )
                    db.session.commit()
                except Exception as e:
                    print(f"Scene Drive update error: {e}")
        drive_bg(_bg, app, scene.id, user.id)
    return jsonify(scene.to_dict())


@app.route('/api/scenes/<int:scene_id>', methods=['DELETE'])
def delete_scene(scene_id):
    scene = Scene.query.get_or_404(scene_id)
    get_project_or_404(scene.project_id)
    drive_file_id = scene.drive_file_id
    db.session.delete(scene)
    db.session.commit()

    user = get_current_user()
    if user and user.access_token and drive_file_id:
        def _bg(app, uid, fid):
            with app.app_context():
                u = User.query.get(uid)
                if not u:
                    return
                delete_drive_file_if_exists(get_drive_service(u), fid)
        drive_bg(_bg, app, user.id, drive_file_id)
    return jsonify({'message': 'Deleted'})


# --- LORE ---

@app.route('/api/projects/<int:project_id>/lore', methods=['GET'])
def get_lore(project_id):
    get_project_or_404(project_id)
    return jsonify([i.to_dict() for i in LoreItem.query.filter_by(project_id=project_id).order_by(LoreItem.name).all()])


@app.route('/api/projects/<int:project_id>/lore', methods=['POST'])
def create_lore(project_id):
    get_project_or_404(project_id)
    data = request.get_json()
    if not data or not data.get('name'):
        return jsonify({'error': 'Name required'}), 400
    item = LoreItem(
        project_id=project_id,
        name=data['name'],
        category=data.get('category', 'item'),
        description=data.get('description', ''),
        image_url=data.get('image_url', ''),
        image_focus=data.get('image_focus', 'center'),
        extra_notes=data.get('extra_notes', ''),
        event_date=data.get('event_date', ''),
        event_order=int(data.get('event_order') or 0),
        show_in_web=bool(data.get('show_in_web', True)),
        web_x=data.get('web_x'),
        web_y=data.get('web_y'),
        map_x=float(data.get('map_x') or 50),
        map_y=float(data.get('map_y') or 50)
    )
    item.set_aliases(data.get('aliases', []))
    db.session.add(item)
    db.session.commit()

    user = get_current_user()
    if user and user.access_token:
        def _bg(app, iid, pid, uid):
            with app.app_context():
                u, p, l = User.query.get(uid), Project.query.get(pid), LoreItem.query.get(iid)
                if not u or not p or not l: return
                try:
                    drive = get_drive_service(u)
                    folders = ensure_project_drive_structure(drive, u, p)
                    folder = folders.get('lore')
                    if not folder:
                        return
                    l.drive_file_id = upsert_drive_text_file(
                        drive, l.name, build_lore_drive_content(l), folder, l.drive_file_id
                    )
                    l.drive_image_file_id = upsert_drive_image_file(
                        drive, l.name, l.image_url, folder, l.drive_image_file_id
                    )
                    db.session.commit()
                except Exception as e:
                    print(f"Lore Drive sync error: {e}")
        drive_bg(_bg, app, item.id, item.project_id, user.id)

    return jsonify(item.to_dict()), 201


@app.route('/api/lore/<int:item_id>', methods=['GET'])
def get_lore_item(item_id):
    item = LoreItem.query.get_or_404(item_id)
    get_project_or_404(item.project_id)
    return jsonify(item.to_dict())


@app.route('/api/lore/<int:item_id>', methods=['PUT'])
def update_lore_item(item_id):
    item = LoreItem.query.get_or_404(item_id)
    get_project_or_404(item.project_id)
    data = request.get_json()
    for f in ['name', 'category', 'description', 'image_url', 'image_focus', 'extra_notes', 'event_date', 'web_x', 'web_y']:
        if f in data:
            setattr(item, f, data[f])
    if 'aliases' in data:
        item.set_aliases(data.get('aliases', []))
    if 'event_order' in data:
        item.event_order = int(data.get('event_order') or 0)
    if 'show_in_web' in data:
        item.show_in_web = bool(data.get('show_in_web'))
    if 'map_x' in data:
        item.map_x = float(data.get('map_x') or 0)
    if 'map_y' in data:
        item.map_y = float(data.get('map_y') or 0)
    db.session.commit()

    user = get_current_user()
    if user and user.access_token:
        def _bg(app, iid, uid):
            with app.app_context():
                u, l = User.query.get(uid), LoreItem.query.get(iid)
                if not u or not l:
                    return
                p = Project.query.get(l.project_id)
                if not p:
                    return
                try:
                    drive = get_drive_service(u)
                    folders = ensure_project_drive_structure(drive, u, p)
                    folder = folders.get('lore')
                    if not folder:
                        return
                    l.drive_file_id = upsert_drive_text_file(
                        drive, l.name, build_lore_drive_content(l), folder, l.drive_file_id
                    )
                    l.drive_image_file_id = upsert_drive_image_file(
                        drive, l.name, l.image_url, folder, l.drive_image_file_id
                    )
                    db.session.commit()
                except Exception as e:
                    print(f"Lore Drive update error: {e}")
        drive_bg(_bg, app, item.id, user.id)
    return jsonify(item.to_dict())


@app.route('/api/lore/<int:item_id>', methods=['DELETE'])
def delete_lore_item(item_id):
    item = LoreItem.query.get_or_404(item_id)
    get_project_or_404(item.project_id)
    drive_file_id = item.drive_file_id
    drive_image_file_id = item.drive_image_file_id
    LoreRelationship.query.filter(
        (LoreRelationship.lore_a_id == item.id) | (LoreRelationship.lore_b_id == item.id)
    ).delete(synchronize_session=False)
    db.session.delete(item)
    db.session.commit()

    user = get_current_user()
    if user and user.access_token:
        def _bg(app, uid, text_fid, image_fid):
            with app.app_context():
                u = User.query.get(uid)
                if not u:
                    return
                drive = get_drive_service(u)
                delete_drive_file_if_exists(drive, text_fid)
                delete_drive_file_if_exists(drive, image_fid)
        drive_bg(_bg, app, user.id, drive_file_id, drive_image_file_id)
    return jsonify({'message': 'Deleted'})


# --- NOTES ---

@app.route('/api/projects/<int:project_id>/notes', methods=['GET'])
def get_notes(project_id):
    get_project_or_404(project_id)
    note = Note.query.filter_by(project_id=project_id).first()
    if not note:
        note = Note(project_id=project_id, content='')
        db.session.add(note)
        db.session.commit()
    return jsonify(note.to_dict())


@app.route('/api/projects/<int:project_id>/notes', methods=['PUT'])
def update_notes(project_id):
    get_project_or_404(project_id)
    note = Note.query.filter_by(project_id=project_id).first()
    if not note:
        note = Note(project_id=project_id, content='')
        db.session.add(note)
    data            = request.get_json()
    note.content    = data.get('content', '')
    note.updated_at = datetime.utcnow()
    db.session.commit()

    user = get_current_user()
    if user and user.access_token:
        def _bg(app, pid, content, uid):
            with app.app_context():
                u, p = User.query.get(uid), Project.query.get(pid)
                if not u or not p: return
                try:
                    drive = get_drive_service(u)
                    if not u.scripvia_folder_id:
                        u.scripvia_folder_id = get_or_create_folder(drive, 'Scripvia')
                        db.session.commit()
                    if not p.drive_folder_id:
                        p.drive_folder_id = get_or_create_folder(drive, p.title, parent_id=u.scripvia_folder_id)
                        db.session.commit()
                    results = drive.files().list(
                        q=f"name='_notes.txt' and '{p.drive_folder_id}' in parents and trashed=false",
                        fields='files(id)').execute()
                    files = results.get('files', [])
                    if files:
                        drive.files().update(
                            fileId=files[0]['id'],
                            media_body=MediaInMemoryUpload(content.encode('utf-8'), mimetype='text/plain')
                        ).execute()
                    else:
                        create_drive_file(drive, '_notes', content, p.drive_folder_id)
                except Exception as e:
                    print(f"Notes Drive sync error: {e}")
        drive_bg(_bg, app, project_id, note.content, user.id)

    return jsonify(note.to_dict())


# --- RELATIONSHIPS ---

@app.route('/api/projects/<int:project_id>/relationships', methods=['GET'])
def get_relationships(project_id):
    get_project_or_404(project_id)
    return jsonify([r.to_dict() for r in get_valid_relationships(project_id)])


@app.route('/api/projects/<int:project_id>/relationships', methods=['POST'])
def create_relationship(project_id):
    get_project_or_404(project_id)
    data = request.get_json()
    if not data or not data.get('char_a_id') or not data.get('char_b_id'):
        return jsonify({'error': 'Both characters required'}), 400
    if data['char_a_id'] == data['char_b_id']:
        return jsonify({'error': 'Cannot relate a character to themselves'}), 400
    characters = Character.query.filter(
        Character.project_id == project_id,
        Character.id.in_([data['char_a_id'], data['char_b_id']])
    ).all()
    if len(characters) != 2:
        return jsonify({'error': 'Both characters must belong to this project'}), 400
    rel = CharacterRelationship(project_id=project_id, char_a_id=data['char_a_id'],
                                char_b_id=data['char_b_id'], relation_type=data.get('relation_type', ''),
                                description=data.get('description', ''), color=data.get('color', '#7b6fb0'))
    db.session.add(rel)
    db.session.commit()

    user = get_current_user()
    if user and user.access_token:
        def _bg(app, pid, uid):
            with app.app_context():
                u = User.query.get(uid)
                p = Project.query.get(pid)
                if not u or not p:
                    return
                try:
                    sync_project_relationships_to_drive(u, p)
                except Exception as e:
                    print(f"Relationship Drive sync error: {e}")
        drive_bg(_bg, app, project_id, user.id)
    return jsonify(rel.to_dict()), 201


@app.route('/api/relationships/<int:rel_id>', methods=['PUT'])
def update_relationship(rel_id):
    rel  = CharacterRelationship.query.get_or_404(rel_id)
    get_project_or_404(rel.project_id)
    data = request.get_json()
    for f in ['relation_type','description','color']:
        if f in data: setattr(rel, f, data[f])
    db.session.commit()

    user = get_current_user()
    if user and user.access_token:
        def _bg(app, rid, uid):
            with app.app_context():
                u = User.query.get(uid)
                r = CharacterRelationship.query.get(rid)
                if not u or not r:
                    return
                p = Project.query.get(r.project_id)
                if not p:
                    return
                try:
                    sync_project_relationships_to_drive(u, p)
                except Exception as e:
                    print(f"Relationship Drive update error: {e}")
        drive_bg(_bg, app, rel.id, user.id)
    return jsonify(rel.to_dict())


@app.route('/api/relationships/<int:rel_id>', methods=['DELETE'])
def delete_relationship(rel_id):
    rel = CharacterRelationship.query.get_or_404(rel_id)
    get_project_or_404(rel.project_id)
    project_id = rel.project_id
    db.session.delete(rel)
    db.session.commit()

    user = get_current_user()
    if user and user.access_token:
        def _bg(app, pid, uid):
            with app.app_context():
                u = User.query.get(uid)
                p = Project.query.get(pid)
                if not u or not p:
                    return
                try:
                    sync_project_relationships_to_drive(u, p)
                except Exception as e:
                    print(f"Relationship Drive delete sync error: {e}")
        drive_bg(_bg, app, project_id, user.id)
    return jsonify({'message': 'Deleted'})


# --- LORE RELATIONSHIPS ---

@app.route('/api/projects/<int:project_id>/lore-relationships', methods=['GET'])
def get_lore_relationships(project_id):
    get_project_or_404(project_id)
    return jsonify([r.to_dict() for r in get_valid_lore_relationships(project_id)])


@app.route('/api/projects/<int:project_id>/lore-relationships', methods=['POST'])
def create_lore_relationship(project_id):
    get_project_or_404(project_id)
    data = request.get_json()
    if not data or not data.get('lore_a_id') or not data.get('lore_b_id'):
        return jsonify({'error': 'Both lore entries required'}), 400
    if data['lore_a_id'] == data['lore_b_id']:
        return jsonify({'error': 'Cannot connect a lore item to itself'}), 400
    lore_items = LoreItem.query.filter(
        LoreItem.project_id == project_id,
        LoreItem.id.in_([data['lore_a_id'], data['lore_b_id']])
    ).all()
    if len(lore_items) != 2:
        return jsonify({'error': 'Both lore entries must belong to this project'}), 400
    rel = LoreRelationship(
        project_id=project_id,
        lore_a_id=data['lore_a_id'],
        lore_b_id=data['lore_b_id'],
        relation_type=data.get('relation_type', 'related to'),
        description=data.get('description', ''),
        color=data.get('color', '#5f8fd6')
    )
    db.session.add(rel)
    db.session.commit()

    user = get_current_user()
    if user and user.access_token:
        def _bg(app, pid, uid):
            with app.app_context():
                u = User.query.get(uid)
                p = Project.query.get(pid)
                if not u or not p:
                    return
                try:
                    sync_project_lore_relationships_to_drive(u, p)
                except Exception as e:
                    print(f"Lore relationship Drive sync error: {e}")
        drive_bg(_bg, app, project_id, user.id)
    return jsonify(rel.to_dict()), 201


@app.route('/api/lore-relationships/<int:rel_id>', methods=['PUT'])
def update_lore_relationship(rel_id):
    rel = LoreRelationship.query.get_or_404(rel_id)
    get_project_or_404(rel.project_id)
    data = request.get_json()
    for f in ['relation_type', 'description', 'color']:
        if f in data:
            setattr(rel, f, data[f])
    db.session.commit()

    user = get_current_user()
    if user and user.access_token:
        def _bg(app, rid, uid):
            with app.app_context():
                u = User.query.get(uid)
                r = LoreRelationship.query.get(rid)
                if not u or not r:
                    return
                p = Project.query.get(r.project_id)
                if not p:
                    return
                try:
                    sync_project_lore_relationships_to_drive(u, p)
                except Exception as e:
                    print(f"Lore relationship Drive update error: {e}")
        drive_bg(_bg, app, rel.id, user.id)
    return jsonify(rel.to_dict())


@app.route('/api/lore-relationships/<int:rel_id>', methods=['DELETE'])
def delete_lore_relationship(rel_id):
    rel = LoreRelationship.query.get_or_404(rel_id)
    get_project_or_404(rel.project_id)
    project_id = rel.project_id
    db.session.delete(rel)
    db.session.commit()

    user = get_current_user()
    if user and user.access_token:
        def _bg(app, pid, uid):
            with app.app_context():
                u = User.query.get(uid)
                p = Project.query.get(pid)
                if not u or not p:
                    return
                try:
                    sync_project_lore_relationships_to_drive(u, p)
                except Exception as e:
                    print(f"Lore relationship Drive delete sync error: {e}")
        drive_bg(_bg, app, project_id, user.id)
    return jsonify({'message': 'Deleted'})


# --- SEARCH ---

@app.route('/api/projects/<int:project_id>/search', methods=['GET'])
def search_project(project_id):
    get_project_or_404(project_id)
    q = request.args.get('q', '').strip().lower()
    if not q or len(q) < 2:
        return jsonify([])

    def snippet(text, query):
        idx = text.lower().find(query)
        if idx < 0: return ''
        s, e = max(0, idx-60), min(len(text), idx+100)
        return ('...' if s > 0 else '') + text[s:e].strip() + ('...' if e < len(text) else '')

    results = []
    for doc in Document.query.filter_by(project_id=project_id).all():
        text = html_to_plain_text(doc.content or '')
        if q in doc.title.lower() or q in text.lower():
            results.append({'type':'chapter','id':doc.id,'title':doc.title,'snippet':snippet(text,q),'icon':'📄'})

    for scene in Scene.query.filter_by(project_id=project_id).all():
        text = html_to_plain_text(scene.content or '')
        if q in scene.title.lower() or q in text.lower():
            results.append({'type':'scene','id':scene.id,'title':scene.title,'snippet':snippet(text,q),'icon':'⚡'})

    for c in Character.query.filter_by(project_id=project_id).all():
        if q in c.name.lower() or q in (c.personality or '').lower() or q in (c.backstory or '').lower():
            results.append({'type':'character','id':c.id,'title':c.name,'snippet':c.role or '','icon':'👤'})

    for l in LoreItem.query.filter_by(project_id=project_id).all():
        if q in l.name.lower() or q in (l.description or '').lower():
            results.append({'type':'lore','id':l.id,'title':l.name,'snippet':l.category or '','icon':'📖'})

    return jsonify(results)


@app.route('/api/projects/<int:project_id>/drive-sync-all', methods=['POST'])
def sync_project_drive_contents(project_id):
    user = get_current_user()
    if not user or not user.access_token:
        return jsonify({'error': 'Not logged in'}), 401
    project = get_project_or_404(project_id)

    def _bg(app, pid, uid):
        with app.app_context():
            u = User.query.get(uid)
            p = Project.query.get(pid)
            if not u or not p:
                return
            try:
                sync_full_project_to_drive(u, p)
            except Exception as e:
                print(f"Project Drive sync error: {e}")

    drive_bg(_bg, app, project.id, user.id)
    return jsonify({'queued': True})


# --- WIKI ---

@app.route('/api/projects/<int:project_id>/wiki', methods=['GET'])
def get_wiki_data(project_id):
    get_project_or_404(project_id)
    chars = Character.query.filter_by(project_id=project_id).all()
    lore  = LoreItem.query.filter_by(project_id=project_id).all()
    wiki  = {}
    for c in chars:
        entry = {
            'type': 'character', 'name': c.name, 'role': c.role, 'age': c.age,
            'titles': c.get_titles(),
            'image_url': c.image_url, 'id': c.id, 'image_focus': c.image_focus,
            'summary':    c.personality[:200] + '...' if len(c.personality) > 200 else c.personality,
            'backstory':  c.backstory[:200]   + '...' if len(c.backstory)   > 200 else c.backstory,
            'appearance': c.appearance[:150]  + '...' if len(c.appearance)  > 150 else c.appearance,
        }
        wiki[c.name.lower()] = entry
        for alias in c.get_titles():
            wiki[alias.lower()] = entry
    for l in lore:
        entry = {
            'type': 'lore', 'name': l.name, 'category': l.category,
            'aliases': l.get_aliases(),
            'image_url': l.image_url, 'id': l.id, 'image_focus': l.image_focus,
            'summary': l.description[:150] + '...' if len(l.description) > 150 else l.description,
        }
        wiki[l.name.lower()] = entry
        for alias in l.get_aliases():
            wiki[alias.lower()] = entry
    return jsonify(wiki)


# --- DRIVE SYNC ---

@app.route('/api/documents/<int:doc_id>/sync', methods=['POST'])
def sync_to_drive(doc_id):
    user = get_current_user()
    if not user or not user.access_token:
        return jsonify({'error': 'Not logged in'}), 401
    doc     = Document.query.get_or_404(doc_id)
    project = get_project_or_404(doc.project_id)
    try:
        had_drive_file = bool(doc.drive_file_id)
        sync_full_project_to_drive(user, project)
        db.session.refresh(doc)
        action = 'updated' if had_drive_file and doc.drive_file_id else 'created'
        return jsonify({
            'success': True,
            'action': action,
            'project_folder_id': project.drive_folder_id,
            'document_drive_file_id': doc.drive_file_id
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/documents/<int:doc_id>/sync-status', methods=['GET'])
def sync_status(doc_id):
    doc = Document.query.get_or_404(doc_id)
    get_project_or_404(doc.project_id)
    return jsonify({'synced': bool(doc.drive_file_id), 'drive_file_id': doc.drive_file_id})


# --- SINGLE DOC EXPORT ---

@app.route('/api/documents/<int:doc_id>/export/pdf', methods=['GET'])
def export_pdf(doc_id):
    from reportlab.lib.enums import TA_LEFT
    doc     = Document.query.get_or_404(doc_id)
    project = get_project_or_404(doc.project_id)
    buffer  = io.BytesIO()
    styles  = getSampleStyleSheet()

    def style(name, parent='Normal', **kw):
        return ParagraphStyle(name, parent=styles[parent], **kw)

    title_s = style('T', 'Title', fontSize=24, textColor=colors.HexColor('#3d3580'), spaceAfter=6, fontName='Helvetica-Bold')
    sub_s   = style('S', fontSize=11, textColor=colors.HexColor('#888888'), spaceAfter=24)
    h1_s    = style('H1', 'Heading1', fontSize=18, textColor=colors.HexColor('#3d3580'), fontName='Helvetica-Bold')
    h2_s    = style('H2', 'Heading2', fontSize=14, textColor=colors.HexColor('#5548a0'), fontName='Helvetica-Bold')
    h3_s    = style('H3', 'Heading3', fontSize=12, textColor=colors.HexColor('#7b6fb0'), fontName='Helvetica-Bold')
    body_s  = style('B', fontSize=11, leading=18, spaceAfter=8, textColor=colors.HexColor('#1a1a2e'))
    quote_s = style('Q', fontSize=11, leading=18, leftIndent=24, fontName='Helvetica-Oblique', textColor=colors.HexColor('#555555'))

    pdf  = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=inch, leftMargin=inch, topMargin=1.2*inch, bottomMargin=inch)
    story = [Paragraph(doc.title, title_s),
             Paragraph(f"from {project.title} · exported via Scripvia", sub_s),
             Spacer(1, 0.2*inch)]

    tag_map = {'h1': h1_s, 'h2': h2_s, 'h3': h3_s}
    for el in BeautifulSoup(doc.content or '', 'html.parser').find_all(['p','h1','h2','h3','blockquote','li']):
        text = el.get_text(strip=True)
        if not text: continue
        if el.name in tag_map:
            story.append(Paragraph(text, tag_map[el.name]))
        elif el.name == 'blockquote':
            story.append(Paragraph(f'"{text}"', quote_s))
        elif el.name == 'li':
            story.append(Paragraph(f'• {text}', body_s))
        else:
            story.append(Paragraph(text, body_s))
        story.append(Spacer(1, 0.05*inch))

    pdf.build(story)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name=f"{doc.title}.pdf", mimetype='application/pdf')


@app.route('/api/documents/<int:doc_id>/export/docx', methods=['GET'])
def export_docx(doc_id):
    doc     = Document.query.get_or_404(doc_id)
    project = get_project_or_404(doc.project_id)
    word    = DocxDocument()
    section = word.sections[0]
    section.top_margin = section.bottom_margin = Inches(1.2)
    section.left_margin = section.right_margin = Inches(1)

    def color(run, h):
        h = h.lstrip('#')
        run.font.color.rgb = RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

    def para(text, size, bold=False, italic=False, hex_color='1a1a2e', indent=None):
        p = word.add_paragraph()
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.font.bold, r.font.italic = bold, italic
        color(r, hex_color)
        if indent: p.paragraph_format.left_indent = indent
        return p

    para(doc.title, 24, bold=True, hex_color='3d3580')
    para(f"from {project.title} · exported via Scripvia", 10, italic=True, hex_color='888888')
    word.add_paragraph()

    tag_cfg = {'h1':(18,True,'3d3580'), 'h2':(14,True,'5548a0'), 'h3':(12,True,'7b6fb0')}
    for el in BeautifulSoup(doc.content or '', 'html.parser').find_all(['p','h1','h2','h3','blockquote','li']):
        text = el.get_text(strip=True)
        if not text: continue
        if el.name in tag_cfg:
            sz, bd, col = tag_cfg[el.name]
            para(text, sz, bold=bd, hex_color=col)
        elif el.name == 'blockquote':
            para(f'"{text}"', 11, italic=True, hex_color='555555', indent=Inches(0.4))
        elif el.name == 'li':
            para(f'• {text}', 11)
        else:
            para(text, 11)

    buffer = io.BytesIO()
    word.save(buffer)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name=f"{doc.title}.docx",
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')


# --- FULL PROJECT EXPORT ---

@app.route('/api/projects/<int:project_id>/export/pdf', methods=['GET'])
def export_project_pdf(project_id):
    project = get_project_or_404(project_id)
    docs    = Document.query.filter_by(project_id=project_id).order_by(
        Document.order_index.asc(), Document.created_at.asc()).all()
    buffer  = io.BytesIO()
    styles  = getSampleStyleSheet()

    def ps(name, parent='Normal', **kw):
        return ParagraphStyle(name, parent=styles[parent], **kw)

    cover_title = ps('CT','Title', fontSize=36, textColor=colors.HexColor('#7b6fb0'), spaceAfter=20, fontName='Times-Bold', alignment=1)
    cover_sub   = ps('CS', fontSize=13, textColor=colors.HexColor('#888888'), spaceAfter=8, alignment=1)
    toc_title   = ps('TT','Heading1', fontSize=20, textColor=colors.HexColor('#7b6fb0'), spaceAfter=20, fontName='Times-Bold')
    toc_entry   = ps('TE', fontSize=12, textColor=colors.HexColor('#333333'), spaceAfter=8, leftIndent=10)
    ch_title    = ps('ChT','Heading1', fontSize=24, textColor=colors.HexColor('#7b6fb0'), spaceAfter=6, fontName='Times-Bold', alignment=1)
    ch_num      = ps('ChN', fontSize=11, textColor=colors.HexColor('#aaaaaa'), spaceAfter=20, alignment=1)
    body        = ps('Bo', fontSize=12, leading=20, textColor=colors.HexColor('#1a1a1a'), spaceAfter=12, firstLineIndent=24)
    h2          = ps('H2','Heading2', fontSize=16, textColor=colors.HexColor('#5548a0'), spaceAfter=10, fontName='Times-Bold')
    h3          = ps('H3','Heading3', fontSize=13, textColor=colors.HexColor('#7b6fb0'), spaceAfter=8,  fontName='Times-Bold')

    total_words = sum(len(html_to_plain_text(d.content or '').split()) for d in docs if d.content)
    els = [Spacer(1, 2*inch), Paragraph(project.title, cover_title), Spacer(1, 0.2*inch)]
    if project.description: els.append(Paragraph(project.description, cover_sub))
    els += [Spacer(1,0.3*inch), Paragraph(f'Genre: {project.genre.title()}', cover_sub),
            Paragraph(f'{len(docs)} chapter{"s" if len(docs)!=1 else ""}', cover_sub),
            Paragraph(f'{total_words:,} words', cover_sub), PageBreak(),
            Paragraph('Table of Contents', toc_title),
            HRFlowable(width='100%', thickness=1, color=colors.HexColor('#cccccc')),
            Spacer(1, 0.15*inch)]
    for i, d in enumerate(docs, 1):
        els.append(Paragraph(f'{i}.  {d.title}', toc_entry))
    els.append(PageBreak())

    for i, doc in enumerate(docs, 1):
        els += [Paragraph(f'Chapter {i}', ch_num), Paragraph(doc.title, ch_title),
                HRFlowable(width='60%', thickness=1, color=colors.HexColor('#7b6fb0'), hAlign='CENTER'),
                Spacer(1, 0.3*inch)]
        if doc.content:
            bq_s = ps('BQ', parent='Normal', leftIndent=30, textColor=colors.HexColor('#666666'), fontName='Times-Italic')
            for el in BeautifulSoup(doc.content, 'html.parser').find_all(['p','h1','h2','h3','h4','blockquote','li']):
                text = el.get_text().strip()
                if not text: continue
                if   el.name == 'h1':                    els.append(Paragraph(text, ch_title))
                elif el.name in ['h2','h3','h4']:        els.append(Paragraph(text, h2 if el.name=='h2' else h3))
                elif el.name == 'blockquote':            els.append(Paragraph(f'"{text}"', bq_s))
                elif el.name == 'li':                    els.append(Paragraph(f'• {text}', body))
                else:                                    els.append(Paragraph(text, body))
        if i < len(docs): els.append(PageBreak())

    SimpleDocTemplate(buffer, pagesize=A4, rightMargin=1.2*inch, leftMargin=1.2*inch,
                      topMargin=1.2*inch, bottomMargin=1.2*inch).build(els)
    buffer.seek(0)
    return send_file(buffer, mimetype='application/pdf', as_attachment=True,
                     download_name=f"{project.title.replace(' ','_')}_complete.pdf")


@app.route('/api/projects/<int:project_id>/export/docx', methods=['GET'])
def export_project_docx(project_id):
    project = get_project_or_404(project_id)
    docs    = Document.query.filter_by(project_id=project_id).order_by(
        Document.order_index.asc(), Document.created_at.asc()).all()
    docx    = DocxDocument()
    section = docx.sections[0]
    section.top_margin = section.bottom_margin = Inches(1.2)
    section.left_margin = section.right_margin = Inches(1.3)

    def color(run, h):
        run.font.color.rgb = RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

    total_words = sum(len(html_to_plain_text(d.content or '').split()) for d in docs if d.content)

    # Cover
    for _ in range(8): docx.add_paragraph('')
    p = docx.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(project.title); r.font.size = Pt(36); r.font.bold = True; color(r, '7b6fb0')
    if project.description:
        p = docx.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(project.description); r.font.size = Pt(13); color(r, '888888')
    p = docx.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'{project.genre.title()} · {len(docs)} chapters · {total_words:,} words')
    r.font.size = Pt(11); color(r, 'aaaaaa')
    docx.add_page_break()

    # TOC
    p = docx.add_paragraph(); r = p.add_run('Table of Contents')
    r.font.size = Pt(20); r.font.bold = True; color(r, '7b6fb0')
    for i, d in enumerate(docs, 1):
        p = docx.add_paragraph(); r = p.add_run(f'{i}.   {d.title}')
        r.font.size = Pt(12); color(r, '333333')
    docx.add_page_break()

    # Chapters
    for i, doc in enumerate(docs, 1):
        p = docx.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f'Chapter {i}'); r.font.size = Pt(11); color(r, 'aaaaaa')
        p = docx.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(doc.title); r.font.size = Pt(24); r.font.bold = True; color(r, '7b6fb0')
        docx.add_paragraph('')

        if doc.content:
            for el in BeautifulSoup(doc.content, 'html.parser').find_all(['p','h1','h2','h3','h4','blockquote','li']):
                text = el.get_text().strip()
                if not text: continue
                p = docx.add_paragraph()
                r = p.add_run(text)
                if   el.name == 'h1':       r.font.size=Pt(20); r.font.bold=True;  color(r,'7b6fb0')
                elif el.name in ['h2','h3','h4']: r.font.size=Pt(16 if el.name=='h2' else 13); r.font.bold=True; color(r,'5548a0')
                elif el.name == 'blockquote':
                    r = p.add_run(f'"{text}"'); r.font.size=Pt(12); r.font.italic=True
                    color(r,'666666'); p.paragraph_format.left_indent=Inches(0.5)
                elif el.name == 'li':       r.font.size=Pt(12)
                else:
                    r.font.size=Pt(12); color(r,'1a1a1a')
                    p.paragraph_format.first_line_indent=Inches(0.3)

        if i < len(docs): docx.add_page_break()

    buffer = io.BytesIO()
    docx.save(buffer)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True,
                     download_name=f"{project.title.replace(' ','_')}_complete.docx",
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')


def ensure_character_titles_column():
    engine = db.engine
    if engine.url.get_backend_name() != 'sqlite':
        return
    with engine.begin() as conn:
        required_columns = {
            'project': {
                'guest_session_id': "VARCHAR(200) DEFAULT ''",
                'map_image_url': "TEXT DEFAULT ''"
            },
            'character': {
                'titles': "TEXT DEFAULT '[]'",
                'drive_file_id': "VARCHAR(200) DEFAULT ''",
                'drive_image_file_id': "VARCHAR(200) DEFAULT ''",
                'web_x': "FLOAT",
                'web_y': "FLOAT"
            },
            'scene': {
                'drive_file_id': "VARCHAR(200) DEFAULT ''"
            },
            'lore_item': {
                'aliases': "TEXT DEFAULT '[]'",
                'drive_file_id': "VARCHAR(200) DEFAULT ''",
                'drive_image_file_id': "VARCHAR(200) DEFAULT ''",
                'event_date': "VARCHAR(120) DEFAULT ''",
                'event_order': "INTEGER DEFAULT 0",
                'show_in_web': "BOOLEAN DEFAULT 1",
                'web_x': "FLOAT",
                'web_y': "FLOAT",
                'map_x': "FLOAT DEFAULT 50",
                'map_y': "FLOAT DEFAULT 50"
            }
        }
        for table_name, columns in required_columns.items():
            cols = conn.exec_driver_sql(f"PRAGMA table_info({table_name})").fetchall()
            col_names = {col[1] for col in cols}
            for column_name, column_def in columns.items():
                if column_name not in col_names:
                    conn.exec_driver_sql(f"ALTER TABLE {table_name} ADD COLUMN {column_name} {column_def}")


def ensure_lore_relationships_table():
    engine = db.engine
    if engine.url.get_backend_name() != 'sqlite':
        return
    with engine.begin() as conn:
        table_names = {row[0] for row in conn.exec_driver_sql("SELECT name FROM sqlite_master WHERE type='table'").fetchall()}
        if 'lore_relationship' not in table_names:
            conn.exec_driver_sql("""
                CREATE TABLE lore_relationship (
                    id INTEGER NOT NULL PRIMARY KEY,
                    project_id INTEGER NOT NULL,
                    lore_a_id INTEGER NOT NULL,
                    lore_b_id INTEGER NOT NULL,
                    relation_type VARCHAR(100) DEFAULT '',
                    description TEXT DEFAULT '',
                    color VARCHAR(20) DEFAULT '#5f8fd6',
                    FOREIGN KEY(project_id) REFERENCES project (id),
                    FOREIGN KEY(lore_a_id) REFERENCES lore_item (id),
                    FOREIGN KEY(lore_b_id) REFERENCES lore_item (id)
                )
            """)


# --- ENTRY POINT ---

if __name__ == '__main__':
    with app.app_context():
        db.create_all()
        ensure_character_titles_column()
        ensure_lore_relationships_table()
        print("Starting app... DB will be created fresh if not exists.")
    print("Scripvia running at http://localhost:5000")
    app.run(debug=False, use_reloader=False)
